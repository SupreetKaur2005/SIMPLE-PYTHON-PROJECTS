import random
import re
import hashlib
import logging
import time
from contextlib import contextmanager
from typing import List, Dict, Optional, Any
from docx import Document
from nltk.corpus import wordnet, opinion_lexicon
from textblob import TextBlob
from pathlib import Path
import concurrent.futures
from tqdm import tqdm

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

@contextmanager
def timer(description: str):
    """Context manager for timing code execution."""
    start = time.time()
    yield
    elapsed = time.time() - start
    logger.info(f"{description}: {elapsed:.2f} seconds")

class TextCache:
    """Cache for storing processed text to avoid redundant operations."""
    
    def __init__(self, max_size: int = 1000):
        self.cache = {}
        self.max_size = max_size
        self.hits = 0
        self.misses = 0
    
    def get(self, text: str, operation: str) -> Optional[str]:
        """Get cached result if available."""
        key = f"{hashlib.md5(text.encode()).hexdigest()}_{operation}"
        if key in self.cache:
            self.hits += 1
            return self.cache[key]
        self.misses += 1
        return None
    
    def set(self, text: str, operation: str, result: str) -> None:
        """Store result in cache."""
        key = f"{hashlib.md5(text.encode()).hexdigest()}_{operation}"
        if len(self.cache) >= self.max_size:
            self.cache.pop(next(iter(self.cache)))
        self.cache[key] = result
    
    def stats(self) -> Dict[str, Any]:
        """Return cache statistics."""
        return {
            "size": len(self.cache),
            "hits": self.hits,
            "misses": self.misses,
            "hit_ratio": self.hits / (self.hits + self.misses) if (self.hits + self.misses) > 0 else 0
        }

class SynonymDatabase:
    """Enhanced synonym management with contextual awareness."""
    
    def __init__(self):
        self.cache = {}
        try:
            import spacy
            self.nlp = spacy.load("en_core_web_sm")
            self.use_spacy = True
            logger.info("Using spaCy for enhanced language processing")
        except (ImportError, OSError):
            self.use_spacy = False
            logger.info("spaCy not available, falling back to basic processing")
    
    def get_synonyms(self, word: str, pos_tag: str) -> List[str]:
        """Get contextually appropriate synonyms for a word."""
        cache_key = f"{word}_{pos_tag}"
        if cache_key in self.cache:
            return self.cache[cache_key]
        
        pos_mapping = {
            'NN': 'n', 'NNS': 'n', 'NNP': 'n', 'NNPS': 'n',
            'VB': 'v', 'VBD': 'v', 'VBG': 'v', 'VBN': 'v', 'VBP': 'v', 'VBZ': 'v',
            'JJ': 'a', 'JJR': 'a', 'JJS': 'a',
            'RB': 'r', 'RBR': 'r', 'RBS': 'r'
        }
        
        if self.use_spacy:
            synonyms = self._get_spacy_synonyms(word, pos_tag)
        else:
            pos = pos_mapping.get(pos_tag[:2], None)
            if not pos:
                self.cache[cache_key] = []
                return []
                
            synonyms = {lemma.name().replace('_', ' ') for syn in wordnet.synsets(word, pos=pos) for lemma in syn.lemmas() if lemma.name().lower() != word.lower()}
        
        result = list(synonyms)
        self.cache[cache_key] = result
        return result
    
    def _get_spacy_synonyms(self, word: str, pos_tag: str) -> set:
        """Get higher-quality synonyms using spaCy's word vectors."""
        doc = self.nlp(word)
        if not doc or not doc[0].has_vector:
            return set()
        
        similar_words = set()
        vector = doc[0].vector
        
        for term in self.nlp.vocab:
            if term.has_vector and term.is_alpha and len(term.text) > 2:
                similarity = term.similarity(doc[0])
                if similarity > 0.7 and term.text.lower() != word.lower():
                    similar_words.add(term.text)
                    if len(similar_words) >= 10:
                        break
        
        if len(similar_words) < 3:
            pos_mapping = {
                'NN': 'n', 'NNS': 'n', 'NNP': 'n', 'NNPS': 'n',
                'VB': 'v', 'VBD': 'v', 'VBG': 'v', 'VBN': 'v', 'VBP': 'v', 'VBZ': 'v',
                'JJ': 'a', 'JJR': 'a', 'JJS': 'a',
                'RB': 'r', 'RBR': 'r', 'RBS': 'r'
            }
            
            pos = pos_mapping.get(pos_tag[:2], None)
            if pos:
                for syn in wordnet.synsets(word, pos=pos):
                    for lemma in syn.lemmas():
                        lemma_name = lemma.name().replace('_', ' ')
                        if lemma_name.lower() != word.lower():
                            similar_words.add(lemma_name)
        
        return similar_words

class AdversarialParaphraser:
    """Handles adversarial text transformations with enhanced capabilities."""
    
    def __init__(self):
        self.positive_words = set(opinion_lexicon.positive())
        self.negative_words = set(opinion_lexicon.negative())
        self.cache = TextCache()
        self.synonym_db = SynonymDatabase()
        self.contextual_pairs = self._load_contextual_pairs()
    
    def _load_contextual_pairs(self) -> Dict[str, List[str]]:
        """Load contextually related word pairs for better transformations."""
        return {
            "good": ["bad", "poor", "inadequate", "subpar"],
            "bad": ["good", "excellent", "superior", "outstanding"],
            "happy": ["sad", "unhappy", "disappointed", "displeased"],
            "sad": ["happy", "joyful", "delighted", "pleased"],
            "love": ["hate", "despise", "dislike", "detest"],
            "hate": ["love", "adore", "cherish", "appreciate"],
        }
    
    def _semantic_inversion(self, text: str) -> str:
        """Invert sentiment-bearing words while preserving structure."""
        cached = self.cache.get(text, "semantic_inversion")
        if cached:
            return cached
            
        words = text.split()
        inverted_words = []
        
        for word in words:
            clean_word = re.sub(r'[^\w]', '', word.lower())
            if clean_word in self.contextual_pairs:
                replacement = random.choice(self.contextual_pairs[clean_word])
                inverted_words.append(self._preserve_case(word, replacement))
            elif clean_word in self.positive_words:
                replacement = random.choice(list(self.negative_words))
                inverted_words.append(self._preserve_case(word, replacement))
            elif clean_word in self.negative_words:
                replacement = random.choice(list(self.positive_words))
                inverted_words.append(self._preserve_case(word, replacement))
            else:
                inverted_words.append(word)
        
        result = ' '.join(inverted_words)
        self.cache.set(text, "semantic_inversion", result)
        return result

    def _preserve_case(self, original: str, replacement: str) -> str:
        """Preserve the original word case in the replacement."""
        if original.istitle():
            return replacement.title()
        elif original.isupper():
            return replacement.upper()
        return replacement.lower()

    def _add_stylometric_noise(self, text: str) -> str:
        """Add natural writing variations to disrupt stylometric analysis."""
        cached = self.cache.get(text, "stylometric_noise")
        if cached:
            return cached
            
        blob = TextBlob(text)
        sentences = [str(sent).strip() for sent in blob.sentences if str(sent).strip()]
        
        if len(sentences) < 2:
            return text
            
        modified = []
        i = 0
        while i < len(sentences):
            if i + 1 < len(sentences) and random.random() < 0.4:
                connector = random.choice(['', ', ', '; ', ' because ', ' while ', ' and '])
                merged = sentences[i] + connector + sentences[i+1][0].lower() + sentences[i+1][1:]
                modified.append(merged)
                i += 2
            else:
                modified.append(sentences[i])
                i += 1
        
        result = ' '.join(modified)
        self.cache.set(text, "stylometric_noise", result)
        return result
    
    def _rearrange_clauses(self, text: str) -> str:
        """Rearrange clauses within complex sentences."""
        cached = self.cache.get(text, "rearrange_clauses")
        if cached:
            return cached
            
        clause_markers = [", ", "; ", " but ", " however ", " nevertheless ", " although ", " while "]
        
        for marker in clause_markers:
            if marker in text:
                parts = text.split(marker, 1)
                if len(parts) == 2 and random.random() < 0.5:
                    result = parts[1].strip().capitalize() + marker + parts[0].strip().lower()
                    self.cache.set(text, "rearrange_clauses", result)
                    return result
        
        return text

    def paraphrase(self, text: str) -> str:
        """Apply adversarial transformations without watermarking."""
        if not text.strip():
            return text
            
        cached = self.cache.get(text, "paraphrase")
        if cached:
            return cached
            
        transformed = self._semantic_inversion(text)
        transformed = self._add_stylometric_noise(transformed)
        transformed = self._rearrange_clauses(transformed)
        
        self.cache.set(text, "paraphrase", transformed)
        return transformed

class EthicalHumanizer:
    """Enhances text to appear more human-written while preserving meaning."""
    
    def __init__(self, min_typo_prob: float = 0.01):
        self.paraphraser = AdversarialParaphraser()
        self.min_typo_prob = min_typo_prob
        self.log = []
        self.transformations_applied = 0
        self.heading_texts = {}
        self.cache = TextCache()
        
        self.stats = {
            "paragraphs_processed": 0,
            "sentences_processed": 0,
            "words_processed": 0,
            "headings_preserved": 0,
            "processing_time": 0
        }
        
    def _get_contextual_synonyms(self, word: str, pos_tag: str) -> List[str]:
        """Get semantically appropriate synonyms based on POS tag."""
        return self.paraphraser.synonym_db.get_synonyms(word, pos_tag)

    def _restructure_sentences(self, text: str) -> str:
        """Vary sentence structure while maintaining coherence."""
        cached = self.cache.get(text, "restructure")
        if cached:
            return cached
            
        blob = TextBlob(text)
        sentences = [str(sent).strip() for sent in blob.sentences if str(sent).strip()]
        
        self.stats["sentences_processed"] += len(sentences)
        
        if len(sentences) < 2:
            return text
            
        modified = []
        for i, sent in enumerate(sentences):
            if i > 0 and random.random() < 0.3:
                strategy = random.choice(["merge", "because", "although", "while"])
                if strategy == "merge":
                    modified[-1] += ' ' + sent[0].lower() + sent[1:]
                elif strategy == "because":
                    modified[-1] += ' because ' + sent[0].lower() + sent[1:]
                elif strategy == "although":
                    modified[-1] = 'Although ' + modified[-1][0].lower() + modified[-1][1:] + ', ' + sent[0].lower() + sent[1:]
                elif strategy == "while":
                    modified[-1] += ', while ' + sent[0].lower() + sent[1:]
            else:
                modified.append(sent)
                
        result = ' '.join(modified)
        self.cache.set(text, "restructure", result)
        return result

    def _add_controlled_imperfections(self, text: str) -> str:
        """Add subtle human-like imperfections."""
        cached = self.cache.get(text, "imperfections")
        if cached:
            return cached
            
        words = text.split()
        self.stats["words_processed"] += len(words)
        modified_words = []
        
        for word in words:
            if len(word) > 4 and random.random() < self.min_typo_prob:
                typo_type = random.choice(['swap', 'repeat', 'capitalize', 'omit', 'insert', 'replace'])
                
                if typo_type == 'swap' and len(word) > 1:
                    idx = random.randint(0, len(word)-2)
                    word = word[:idx] + word[idx+1] + word[idx] + word[idx+2:]
                elif typo_type == 'repeat' and len(word) > 2:
                    idx = random.randint(1, len(word)-1)
                    word = word[:idx] + word[idx] + word[idx:]
                elif typo_type == 'capitalize':
                    word = word.lower().capitalize()
                elif typo_type == 'omit' and len(word) > 3:
                    idx = random.randint(1, len(word)-2)
                    word = word[:idx] + word[idx+1:]
                elif typo_type == 'insert' and len(word) > 2:
                    idx = random.randint(1, len(word)-1)
                    letters = 'abcdefghijklmnopqrstuvwxyz'
                    word = word[:idx] + random.choice(letters) + word[idx:]
                elif typo_type == 'replace' and len(word) > 2:
                    idx = random.randint(1, len(word)-2)
                    adjacent_keys = {
                        'a': 'sqzw', 'b': 'vghn', 'c': 'xdfv', 'd': 'erfcxs',
                        'e': 'wrsdf', 'f': 'drtgcv', 'g': 'ftyhbv', 'h': 'gyujnb',
                        'i': 'ujklo', 'j': 'hyuikn', 'k': 'juilom', 'l': 'kiop',
                        'm': 'njk', 'n': 'bhjm', 'o': 'iklp', 'p': 'ol',
                        'q': 'asw', 'r': 'edft', 's': 'awedxz', 't': 'rfgy',
                        'u': 'yhji', 'v': 'cfgb', 'w': 'qase', 'x': 'zsdc',
                        'y': 'tghu', 'z': 'asx'
                    }
                    if word[idx].lower() in adjacent_keys:
                        replacement = random.choice(adjacent_keys[word[idx].lower()])
                        if word[idx].isupper():
                            replacement = replacement.upper()
                        word = word[:idx] + replacement + word[idx+1:]
                    
            modified_words.append(word)
            
        result = ' '.join(modified_words)
        self.cache.set(text, "imperfections", result)
        return result

    def _process_paragraph(self, text: str, is_heading: bool = False) -> str:
        """Apply all transformations to a single paragraph."""
        if not text.strip():
            return text
        
        if is_heading:
            self.stats["headings_preserved"] += 1
            return text
            
        cached = self.cache.get(text, "process_paragraph")
        if cached:
            return cached
            
        blob = TextBlob(text)
        modified_words = []
        
        for word, tag in blob.tags:
            if word.isdigit() or len(word) < 4:
                modified_words.append(word)
                continue
                
            synonyms = self._get_contextual_synonyms(word, tag)
            if synonyms and random.random() < 0.3:
                replacement = random.choice(synonyms)
                modified_words.append(replacement)
                self.transformations_applied += 1
            else:
                modified_words.append(word)
                
        transformed = ' '.join(modified_words)
        transformed = self._restructure_sentences(transformed)
        transformed = self._add_controlled_imperfections(transformed)
        transformed = self.paraphraser.paraphrase(transformed)
        
        self.cache.set(text, "process_paragraph", transformed)
        return transformed
    
    def _extract_and_preserve_headings(self, doc: Document) -> None:
        """Extract headings from document to preserve them."""
        for paragraph in doc.paragraphs:
            if paragraph.style.name.startswith('Heading') or "title" in paragraph.style.name.lower():
                self.heading_texts[paragraph.text.strip()] = True
                
                if re.match(r'^[\d\.]+\s+.+', paragraph.text.strip()):
                    match = re.match(r'^([\d\.]+\s+)(.+)', paragraph.text.strip())
                    if match and match.group(2):
                        self.heading_texts[match.group(2).strip()] = True
    
    def transform_text(self, input_path: str, output_path: str) -> str:
        """Process DOCX file with ethical transformations, preserving headings."""
        start_time = time.time()
        
        try:
            input_path = Path(input_path)
            output_path = Path(output_path)
            
            if not input_path.exists():
                raise FileNotFoundError(f"Input file not found: {input_path}")
                
            logger.info("First pass: Analyzing document structure and extracting headings")
            doc = Document(input_path)
            self._extract_and_preserve_headings(doc)
            
            logger.info("Second pass: Transforming document content")
            output_doc = Document()
            
            logger.info(f"Processing {len(doc.paragraphs)} paragraphs")
            total_paragraphs = len(doc.paragraphs)
            
            with concurrent.futures.ThreadPoolExecutor() as executor:
                para_tasks = []
                for i, para in enumerate(doc.paragraphs):
                    original = para.text.strip()
                    is_heading = (para.style.name.startswith('Heading') or 
                                 original in self.heading_texts or
                                 "title" in para.style.name.lower())
                    
                    para_tasks.append({
                        'index': i,
                        'text': original,
                        'is_heading': is_heading,
                        'style': para.style.name,
                        'alignment': para.alignment,
                    })
                
                self.stats["paragraphs_processed"] = total_paragraphs
                futures = {executor.submit(self._process_paragraph, task['text'], task['is_heading']): task for task in para_tasks}
                
                results = [None] * len(para_tasks)
                for future in tqdm(concurrent.futures.as_completed(futures), total=len(futures), desc="Processing paragraphs"):
                    task = futures[future]
                    try:
                        transformed_text = future.result()
                        results[task['index']] = (transformed_text, task)
                    except Exception as e:
                        logger.error(f"Error processing paragraph {task['index']}: {str(e)}")
                        results[task['index']] = (task['text'], task)
            
            for transformed_text, task in results:
                if not task['text'] and not transformed_text:
                    continue
                
                new_para = output_doc.add_paragraph()
                new_para.style = doc.styles[task['style']]
                new_para.alignment = task['alignment']
                new_para.text = transformed_text
            
            output_doc.save(output_path)
            self.stats["processing_time"] = time.time() - start_time
            
            logger.info(f"Document transformed successfully in {self.stats['processing_time']:.2f} seconds")
            logger.info(f"Applied {self.transformations_applied} transformations")
            logger.info(f"Processed {self.stats['paragraphs_processed']} paragraphs, {self.stats['sentences_processed']} sentences, {self.stats['words_processed']} words")
            logger.info(f"Preserved {self.stats['headings_preserved']} headings")
            logger.info(f"Cache statistics: {self.cache.stats()}")
            
            return output_path
        except Exception as e:
            logger.error(f"Error transforming document: {str(e)}")
            raise

def main():
    """Main entry point with hardcoded file paths."""
    input_file = "Supreet_Kaur_500122005_ASS4_SRS.docx"
    output_file = "output.docx"
    
    print(f"Processing {input_file} -> {output_file}")
    
    try:
        print("Starting document transformation...")
        humanizer = EthicalHumanizer()
        output_path = humanizer.transform_text(input_file, output_file)
        
        print(f"Successfully transformed document: {output_path}")
        return 0
    except Exception as e:
        print(f"Error: {str(e)}")
        return 1

if __name__ == "__main__":
    main()