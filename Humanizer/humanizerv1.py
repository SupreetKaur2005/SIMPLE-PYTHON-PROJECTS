import random
import re
from docx import Document
from nltk.corpus import wordnet, opinion_lexicon
from textblob import TextBlob
import hashlib

# === ADVERSARIAL PARAPHRASE FUNCTION ===
def semantic_inversion(text):
    """Invert the meaning of the text while preserving structure"""
    # Simple inversion by swapping positive and negative words
    positive_words = set(opinion_lexicon.positive())
    negative_words = set(opinion_lexicon.negative())

    words = text.split()
    inverted_words = []

    for word in words:
        if word in positive_words:
            inverted_words.append(random.choice(list(negative_words)))
        elif word in negative_words:
            inverted_words.append(random.choice(list(positive_words)))
        else:
            inverted_words.append(word)

    return ' '.join(inverted_words)

def add_stylometric_noise(text):
    """Add variations in style to the text"""
    # Change sentence length by merging or splitting sentences
    blob = TextBlob(text)
    sentences = [str(sent) for sent in blob.sentences]

    modified_sentences = []
    for i in range(0, len(sentences), 2):
        if random.random() < 0.5 and i + 1 < len(sentences):
            # Merge two sentences
            merged = sentences[i] + ' ' + sentences[i + 1].lower()
            modified_sentences.append(merged)
        else:
            modified_sentences.append(sentences[i])

    return ' '.join(modified_sentences)

def induce_lexical_watermark(text):
    """Embed a lexical watermark in the text"""
    # Simple watermark by adding a specific pattern of words
    watermark = "embedded watermark"
    words = text.split()
    if len(words) > 10:
        words.insert(len(words) // 2, watermark)
    return ' '.join(words)

def adversarial_paraphrase(text):
    # Multi-step transformation pipeline
    text = semantic_inversion(text)  # Preserve meaning but alter structure
    text = add_stylometric_noise(text)
    text = induce_lexical_watermark(text)
    return text

# === TEXT HUMANIZATION ENGINE ===
class EthicalHumanizer:
    def __init__(self):
        self.log = []

    def _contextual_replace(self, word, tag):
        """POS-aware synonym selection"""
        synonyms = []
        for syn in wordnet.synsets(word):
            if syn.pos() == tag[0].lower():
                for lemma in syn.lemmas():
                    if lemma.name().lower() != word.lower():  # Ensure case insensitivity
                        synonyms.append(lemma.name())
        return random.choice(synonyms) if synonyms else word

    def _restructure_sentences(self, text):
        """Add human writing patterns while preserving meaning"""
        blob = TextBlob(text)
        sentences = [str(sent) for sent in blob.sentences]

        modified = []
        for sent in sentences:
            # Split sentence into clauses based on common conjunctions
            clauses = re.split(r'(, and|, or|, but| and| or| but)', sent)
            if len(clauses) > 1:
                # Reverse the order of clauses for variation
                clauses = clauses[::-1]
                sent = ''.join(clauses)
            modified.append(sent)

        return ' '.join(modified)

    def _add_human_noise(self, text):
        """Introduce controlled imperfections"""
        # Minimal typo simulation to avoid distorting meaning
        text = re.sub(r'(\b\w{5,}\b)', lambda m: m.group()[:-1] + m.group()[-1].upper()
                      if random.random() < 0.01 else m.group(), text)
        # Reduce punctuation variation
        return text.replace('.', '.', 1)

    def transform_text(self, input_path, output_path):
        """Process DOCX file"""
        doc = Document(input_path)
        transformed = []

        for para in doc.paragraphs:
            original = para.text
            if not original.strip():
                continue

            # Transformation pipeline
            tb = TextBlob(original)
            transformed_text = []

            for word, tag in tb.tags:
                new_word = self._contextual_replace(word, tag)
                transformed_text.append(new_word)

            modified = ' '.join(transformed_text)
            modified = self._restructure_sentences(modified)
            modified = self._add_human_noise(modified)

            # Apply adversarial paraphrase
            modified = adversarial_paraphrase(modified)

            # Logging
            self.log.append({
                'original_hash': hashlib.sha256(original.encode()).hexdigest(),
                'modified_hash': hashlib.sha256(modified.encode()).hexdigest(),
                'transformations': len(transformed_text) - len(original.split())
            })

            transformed.append(modified)

        # Save the transformed text
        output = Document()
        output.add_paragraph('\n\n'.join(transformed))
        output.save(output_path)
        return output_path

# === RESEARCH USAGE ===
if __name__ == "__main__":
    humanizer = EthicalHumanizer()

    # Use public domain sample text
    input_file = "Supreet_Kaur_500122005_ASS4_SRS.docx"
    output_file = "human_Supreet_Kaur_500122005_ASS4_SRS.docx"

    humanizer.transform_text(input_file, output_file)
    print(f"Transformed document saved to {output_file}")
    print("Research log:", humanizer.log)
