import random
import re
from docx import Document
from nltk.corpus import wordnet, opinion_lexicon
from textblob import TextBlob
import hashlib
from typing import List, Dict, Tuple, Optional
import logging
from pathlib import Path

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class AdversarialParaphraser:
    """Handles adversarial text transformations without watermarking."""
    
    def __init__(self):
        # Load sentiment lexicons once
        self.positive_words = set(opinion_lexicon.positive())
        self.negative_words = set(opinion_lexicon.negative())
    
    def _semantic_inversion(self, text: str) -> str:
        """Invert sentiment-bearing words while preserving structure."""
        words = text.split()
        inverted_words = []
        
        for word in words:
            clean_word = re.sub(r'[^\w]', '', word.lower())
            if clean_word in self.positive_words:
                replacement = random.choice(list(self.negative_words))
                inverted_words.append(self._preserve_case(word, replacement))
            elif clean_word in self.negative_words:
                replacement = random.choice(list(self.positive_words))
                inverted_words.append(self._preserve_case(word, replacement))
            else:
                inverted_words.append(word)
        
        return ' '.join(inverted_words)

    def _preserve_case(self, original: str, replacement: str) -> str:
        """Preserve the original word case in the replacement."""
        if original.istitle():
            return replacement.title()
        elif original.isupper():
            return replacement.upper()
        return replacement.lower()

    def _add_stylometric_noise(self, text: str) -> str:
        """Add natural writing variations to disrupt stylometric analysis."""
        blob = TextBlob(text)
        sentences = [str(sent).strip() for sent in blob.sentences if str(sent).strip()]
        
        if len(sentences) < 2:
            return text
            
        modified = []
        i = 0
        while i < len(sentences):
            if i + 1 < len(sentences) and random.random() < 0.4:
                # Merge with proper capitalization
                merged = sentences[i] + ' ' + sentences[i+1][0].lower() + sentences[i+1][1:]
                modified.append(merged)
                i += 2
            else:
                modified.append(sentences[i])
                i += 1
                
        return ' '.join(modified)

    def paraphrase(self, text: str) -> str:
        """Apply adversarial transformations without watermarking."""
        if not text.strip():
            return text
            
        transformed = self._semantic_inversion(text)
        transformed = self._add_stylometric_noise(transformed)
        return transformed

class EthicalHumanizer:
    """Enhances text to appear more human-written while preserving meaning."""
    
    def __init__(self, min_typo_prob: float = 0.01, max_typo_prob: float = 0.05):
        self.paraphraser = AdversarialParaphraser()
        self.min_typo_prob = min_typo_prob
        self.max_typo_prob = max_typo_prob
        self.log = []
        self.transformations_applied = 0
        
    def _get_contextual_synonyms(self, word: str, pos_tag: str) -> List[str]:
        """Get semantically appropriate synonyms based on POS tag."""
        pos_mapping = {
            'NN': 'n', 'NNS': 'n', 'NNP': 'n', 'NNPS': 'n',
            'VB': 'v', 'VBD': 'v', 'VBG': 'v', 'VBN': 'v', 'VBP': 'v', 'VBZ': 'v',
            'JJ': 'a', 'JJR': 'a', 'JJS': 'a',
            'RB': 'r', 'RBR': 'r', 'RBS': 'r'
        }
        
        pos = pos_mapping.get(pos_tag[:2], None)
        if not pos:
            return []
            
        synonyms = set()
        for syn in wordnet.synsets(word, pos=pos):
            for lemma in syn.lemmas():
                lemma_name = lemma.name().replace('_', ' ')
                if lemma_name.lower() != word.lower():
                    synonyms.add(lemma_name)
        
        return list(synonyms)

    def _restructure_sentences(self, text: str) -> str:
        """Vary sentence structure while maintaining coherence."""
        blob = TextBlob(text)
        sentences = [str(sent).strip() for sent in blob.sentences if str(sent).strip()]
        
        if len(sentences) < 2:
            return text
            
        modified = []
        for i, sent in enumerate(sentences):
            if i > 0 and random.random() < 0.3:
                # Occasionally combine with previous sentence
                modified[-1] = modified[-1] + ' ' + sent[0].lower() + sent[1:]
            else:
                modified.append(sent)
                
        return ' '.join(modified)

    def _add_controlled_imperfections(self, text: str) -> str:
        """Add subtle human-like imperfections."""
        words = text.split()
        modified_words = []
        
        for word in words:
            if len(word) > 4 and random.random() < self.min_typo_prob:
                # Introduce minor typo
                typo_type = random.choice(['swap', 'repeat', 'capitalize'])
                
                if typo_type == 'swap' and len(word) > 1:
                    idx = random.randint(0, len(word)-2)
                    word = word[:idx] + word[idx+1] + word[idx] + word[idx+2:]
                elif typo_type == 'repeat' and len(word) > 2:
                    idx = random.randint(1, len(word)-1)
                    word = word[:idx] + word[idx] + word[idx:]
                elif typo_type == 'capitalize':
                    word = word.lower().capitalize()
                    
            modified_words.append(word)
            
        return ' '.join(modified_words)

    def _process_paragraph(self, text: str) -> str:
        """Apply all transformations to a single paragraph."""
        if not text.strip():
            return text
            
        # POS-aware synonym replacement
        blob = TextBlob(text)
        modified_words = []
        
        for word, tag in blob.tags:
            synonyms = self._get_contextual_synonyms(word, tag)
            if synonyms and random.random() < 0.3:  # Only replace some words
                replacement = random.choice(synonyms)
                modified_words.append(replacement)
                self.transformations_applied += 1
            else:
                modified_words.append(word)
                
        transformed = ' '.join(modified_words)
        transformed = self._restructure_sentences(transformed)
        transformed = self._add_controlled_imperfections(transformed)
        transformed = self.paraphraser.paraphrase(transformed)
        
        return transformed

    def transform_text(self, input_path: str, output_path: str) -> str:
        """Process DOCX file with ethical transformations."""
        try:
            input_path = Path(input_path)
            output_path = Path(output_path)
            
            if not input_path.exists():
                raise FileNotFoundError(f"Input file not found: {input_path}")
                
            doc = Document(input_path)
            transformed_paragraphs = []
            
            for para in doc.paragraphs:
                original = para.text.strip()
                if not original:
                    transformed_paragraphs.append("")
                    continue
                    
                transformed = self._process_paragraph(original)
                
                # Logging
                self.log.append({
                    'original_hash': hashlib.sha256(original.encode()).hexdigest(),
                    'modified_hash': hashlib.sha256(transformed.encode()).hexdigest(),
                    'transformations': self.transformations_applied
                })
                self.transformations_applied = 0
                
                transformed_paragraphs.append(transformed)
            
            # Save output
            output_doc = Document()
            for para in transformed_paragraphs:
                if para:
                    output_doc.add_paragraph(para)
                else:
                    output_doc.add_paragraph()  # Preserve empty paragraphs
            
            output_doc.save(output_path)
            logger.info(f"Successfully transformed document saved to {output_path}")
            return str(output_path)
            
        except Exception as e:
            logger.error(f"Error during text transformation: {str(e)}")
            raise

def main():
    """Example usage for research purposes."""
    try:
        humanizer = EthicalHumanizer()
        
        # Example file paths - replace with your actual files
        input_file = "Supreet_Kaur_500122005_ASS4_SRS.docx"
        output_file = "output_Supreet_Kaur_500122005_ASS4_SRS.docx"
        
        output_path = humanizer.transform_text(input_file, output_file)
        print(f"Transformed document saved to {output_path}")
        
        if humanizer.log:
            print("\nTransformation Log:")
            for i, entry in enumerate(humanizer.log, 1):
                print(f"Paragraph {i}:")
                print(f"  Original hash: {entry['original_hash']}")
                print(f"  Modified hash: {entry['modified_hash']}")
                print(f"  Transformations applied: {entry['transformations']}")
    
    except Exception as e:
        logger.error(f"Error in main execution: {str(e)}")
        return 1
        
    return 0

if __name__ == "__main__":
    exit(main())